"""Knowledge Base API 完整测试"""
import requests
import json
import sys

BASE = "http://localhost:8001/api"

def test(name, fn):
    try:
        fn()
        print(f"  PASS: {name}")
    except Exception as e:
        print(f"  FAIL: {name} — {e}")

def main():
    # ── Login ──
    print("===== 1. 登录 =====")
    r = requests.post(f"{BASE}/auth/login", json={"username": "admin", "password": "admin123"})
    assert r.status_code == 200, f"login failed: {r.status_code}"
    token = r.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    print(f"  Token: {token[:40]}...")

    # ── Upload Marketing Knowledge (txt) ──
    print("\n===== 2. 上传市场部知识文件 =====")
    txt_content = """这是AI管理系统知识库测试文档。
知识点1：Flutter是Google开源的跨平台UI框架，使用Dart语言开发。
知识点2：AI管理系统包含五大业务模块——市场部、招投标、项目管理、HR、财务。
知识点3：市场部功能包括客户管理、方案生成、项目跟进、社群运营、知识库管理。
知识点4：招投标功能包括合同中心、知识库、招投标流程管理、供应商和讲师管理。"""

    files = {"file": ("test_knowledge.txt", txt_content.encode("utf-8"), "text/plain")}
    data = {"tags": "测试,Flutter,知识库"}
    r = requests.post(f"{BASE}/marketing/knowledge/upload", files=files, data=data, headers=headers, timeout=30)
    assert r.status_code == 200, f"upload failed: {r.status_code} {r.text}"
    entry = r.json()
    print(f"  ID: {entry['id']}")
    print(f"  Title: {entry['title']}")
    print(f"  Content length: {len(entry['content'])} chars")
    print(f"  File ID: {entry.get('source_file_id', 'N/A')}")
    entry_id = entry["id"]

    # ── Upload Marketing Knowledge (PDF) ──
    print("\n===== 3. 上传市场部 PDF 文件 =====")
    # Create a minimal PDF
    pdf_bytes = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>>>endobj
xref
0 4
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
trailer<</Size 4/Root 1 0 R>>
startxref
206
%%EOF"""
    files = {"file": ("test_report.pdf", pdf_bytes, "application/pdf")}
    data = {"tags": "测试,PDF"}
    r = requests.post(f"{BASE}/marketing/knowledge/upload", files=files, data=data, headers=headers, timeout=30)
    assert r.status_code == 200, f"PDF upload failed: {r.status_code} {r.text}"
    pdf_entry = r.json()
    print(f"  PDF uploaded: {pdf_entry['title']}")
    print(f"  Content length: {len(pdf_entry['content'])} chars")

    # ── Upload Word file ──
    print("\n===== 4. 上传市场部 Word 文件 =====")
    from docx import Document
    from io import BytesIO
    doc = Document()
    doc.add_heading("市场部工作规范", 0)
    doc.add_paragraph("市场部负责公司的市场推广、客户管理、品牌建设等工作。")
    doc.add_paragraph("客户满意度是市场部核心考核指标之一。")
    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    files = {"file": ("marketing_guide.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = requests.post(f"{BASE}/marketing/knowledge/upload", files=files, data={"tags": "规范,Word"}, headers=headers, timeout=30)
    assert r.status_code == 200, f"docx upload failed: {r.status_code} {r.text}"
    docx_entry = r.json()
    print(f"  Word uploaded: {docx_entry['title']}")
    print(f"  Content length: {len(docx_entry['content'])} chars")

    # ── Knowledge List ──
    print("\n===== 5. 市场部知识库列表 =====")
    r = requests.get(f"{BASE}/marketing/knowledge", params={"limit": 10}, headers=headers, timeout=10)
    assert r.status_code == 200
    items = r.json()["items"]
    print(f"  条目数: {len(items)}")
    for item in items:
        has_file = "yes" if item.get("source_file_id") else "no"
        print(f"    - {item['title']} (file: {has_file}, preview: {item.get('content_preview', '')[:50]}...)")

    # ── QA: First question ──
    print("\n===== 6. RAG 问答 — 第1轮 =====")
    r = requests.post(f"{BASE}/marketing/knowledge/qa", json={
        "question": "AI管理系统有哪些模块？",
        "top_k": 3,
        "history": []
    }, headers=headers, timeout=120)
    assert r.status_code == 200
    qa1 = r.json()
    print(f"  Q: AI管理系统有哪些模块？")
    print(f"  A: {qa1['answer'][:200]}...")
    print(f"  Model: {qa1['model']}, Sources: {len(qa1['sources'])}")
    if qa1['sources']:
        for s in qa1['sources']:
            print(f"    - [{s.get('title')}]: {s.get('content_preview', '')[:60]}...")

    # ── QA: Follow-up with history ──
    print("\n===== 7. 多轮对话 — 第2轮 =====")
    r = requests.post(f"{BASE}/marketing/knowledge/qa", json={
        "question": "市场部具体有哪些功能？",
        "top_k": 3,
        "history": [
            {"role": "user", "content": "AI管理系统有哪些模块？"},
            {"role": "assistant", "content": qa1["answer"]}
        ]
    }, headers=headers, timeout=120)
    assert r.status_code == 200, f"multi-turn QA failed: {r.status_code} {r.text}"
    qa2 = r.json()
    print(f"  Q: 市场部具体有哪些功能？")
    print(f"  A: {qa2['answer'][:200]}...")
    print(f"  Model: {qa2['model']}, Sources: {len(qa2['sources'])}")

    # ── File Download URL ──
    print("\n===== 8. 文件下载链接 =====")
    r = requests.get(f"{BASE}/marketing/knowledge/{entry_id}/file-url", headers=headers, timeout=10)
    assert r.status_code == 200
    file_info = r.json()
    print(f"  Filename: {file_info['name']}")
    print(f"  Size: {file_info['size_bytes']} bytes")
    print(f"  URL available: {'yes' if file_info.get('url') else 'no'}")

    # ── Upload Bidding Knowledge ──
    print("\n===== 9. 上传招投标知识文件 =====")
    bid_content = """招投标全流程管理文档。
第一阶段：招标准备——市场调研、招标方案编制、资金落实。
第二阶段：招标公告——发布招标公告、资格预审。
第三阶段：投标——投标人编制投标文件、递交投标保证金。
第四阶段：开标评标——开标会议、专家评审、技术商务评分。
第五阶段：定标签约——确定中标人、发放中标通知书、签订合同。"""
    files = {"file": ("bidding_process.txt", bid_content.encode("utf-8"), "text/plain")}
    data = {"tags": "招投标,流程,规范"}
    r = requests.post(f"{BASE}/bidding/knowledge/docs/upload", files=files, data=data, headers=headers, timeout=30)
    assert r.status_code == 200, f"bidding upload failed: {r.status_code} {r.text}"
    bid_entry = r.json()
    print(f"  ID: {bid_entry['id']}")
    print(f"  Content length: {len(bid_entry['content'])} chars")

    # ── Bidding QA ──
    print("\n===== 10. 招投标知识库 QA =====")
    r = requests.post(f"{BASE}/bidding/knowledge/qa", json={
        "question": "招投标有哪些阶段？",
        "top_k": 3,
        "history": []
    }, headers=headers, timeout=120)
    assert r.status_code == 200
    bqa = r.json()
    print(f"  Q: 招投标有哪些阶段？")
    print(f"  A: {bqa['answer'][:200]}...")
    print(f"  Model: {bqa['model']}, Sources: {len(bqa['sources'])}")

    # ── Delete file ──
    print("\n===== 11. 清除文件 =====")
    r = requests.delete(f"{BASE}/marketing/knowledge/{docx_entry['id']}/file", headers=headers, timeout=10)
    assert r.status_code == 200
    print(f"  Word file cleared: {r.json()['message']}")

    # ── Delete entry ──
    print("\n===== 12. 删除知识条目 =====")
    r = requests.delete(f"{BASE}/marketing/knowledge/{entry_id}", headers=headers, timeout=10)
    assert r.status_code == 200
    print(f"  Entry deleted: {r.json()['message']}")

    print("\n===== 全部 12 项测试通过! =====")

if __name__ == "__main__":
    main()
